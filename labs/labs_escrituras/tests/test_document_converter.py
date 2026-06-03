from __future__ import annotations

import subprocess
from pathlib import Path

import pytest
from docx import Document

from lab_escrituras import document_converter as dc
from lab_escrituras.config import LabConfig
from lab_escrituras.pdf_inspector import PdfInspectionResult


def config(tmp_path, *, ocr_enabled=False) -> LabConfig:
    return LabConfig(
        db_url="postgresql://example",
        supabase_url=None,
        supabase_service_role_key=None,
        pdf_inspector_command="pdf2md",
        output_dir=tmp_path / "output",
        ocr_enabled=ocr_enabled,
        ocr_language="spa",
        ocr_dpi=220,
        openai_api_key=None,
        embedding_model="text-embedding-3-small",
        embedding_batch_size=64,
    )


class FakePdfPage:
    def __init__(self, text: str):
        self._text = text

    def extract_text(self) -> str:
        return self._text


class FakePdfReader:
    def __init__(self, _path: str):
        self.pages = [FakePdfPage("Primera pagina"), FakePdfPage("Segunda pagina")]


def test_convert_pdf_persists_physical_pages(monkeypatch, tmp_path):
    monkeypatch.setattr(dc, "PdfReader", FakePdfReader)
    monkeypatch.setattr(
        dc,
        "inspect_pdf",
        lambda *_args: PdfInspectionResult(
            pdf_type="TextBased",
            confidence=0.95,
            page_count=2,
            markdown="",
            raw={"engine": "test"},
        ),
    )

    result = dc.convert_pdf(config(tmp_path), tmp_path / "document.pdf")

    assert result.source_format == "pdf"
    assert result.page_count == 2
    assert [page.page_number for page in result.pages] == [1, 2]
    assert "Primera pagina" in result.pages[0].markdown
    assert "Segunda pagina" in result.pages[1].markdown


def test_convert_docx_extracts_paragraphs_and_tables(tmp_path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Escritura", level=1)
    document.add_paragraph("Comparecen las partes.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Campo"
    table.rows[0].cells[1].text = "Valor"
    table.rows[1].cells[0].text = "Lote"
    table.rows[1].cells[1].text = "12"
    document.save(path)

    result = dc.convert_docx(path)

    markdown = "\n\n".join(page.markdown for page in result.pages)
    assert result.source_format == "docx"
    assert "# Escritura" in markdown
    assert "Comparecen las partes." in markdown
    assert "| Campo | Valor |" in markdown


def test_convert_textutil_document_uses_logical_pages(monkeypatch, tmp_path):
    path = tmp_path / "sample.rtf"
    path.write_text("{\\rtf1 ejemplo}", encoding="utf-8")

    def fake_run(args, **_kwargs):
        output_path = Path(args[4])
        output_path.write_text("Pagina uno\fPagina dos", encoding="utf-8")
        return subprocess.CompletedProcess(args, 0)

    monkeypatch.setattr(dc, "_textutil_binary", lambda: "/usr/bin/textutil")
    monkeypatch.setattr(dc.subprocess, "run", fake_run)

    result = dc.convert_textutil_document(path, "rtf")

    assert result.source_format == "rtf"
    assert [page.page_number for page in result.pages] == [1, 2]
    assert "Pagina uno" in result.pages[0].markdown
    assert "Pagina dos" in result.pages[1].markdown


def test_textutil_missing_is_explicit(monkeypatch):
    monkeypatch.setattr(dc.shutil, "which", lambda _name: None)
    monkeypatch.setattr(dc.Path, "exists", lambda _self: False)

    with pytest.raises(RuntimeError, match="textutil is required"):
        dc._textutil_binary()
