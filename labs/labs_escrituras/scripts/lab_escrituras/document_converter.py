from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field, replace
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from .config import LabConfig
from .ocr import ocr_pdf_pages
from .pdf_inspector import inspect_pdf, should_require_ocr


@dataclass(frozen=True)
class ConvertedPage:
    page_number: int
    markdown: str
    needs_ocr: bool = False
    has_encoding_issues: bool = False
    is_complex_layout: bool = False
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ConvertedDocument:
    source_format: str
    detected_type: str
    confidence: float
    page_count: int
    pages: list[ConvertedPage]
    pages_needing_ocr: list[int] = field(default_factory=list)
    has_encoding_issues: bool = False
    is_complex_layout: bool = False
    raw: dict[str, Any] = field(default_factory=dict)


def _clean_text(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.splitlines()).strip()


def _page_to_markdown(page_number: int, text: str) -> str:
    cleaned = _clean_text(text)
    if not cleaned:
        return f"<!-- Page {page_number}: no extractable text -->"
    return f"<!-- Page {page_number} -->\n\n{cleaned}"


def _apply_pdf_ocr_fallback(config: LabConfig, pdf_path: Path, converted: ConvertedDocument) -> ConvertedDocument:
    if not config.ocr_enabled or not converted.pages_needing_ocr:
        return converted

    pages_by_number = {page.page_number: page for page in converted.pages}
    updated_pages: list[ConvertedPage] = []
    unresolved_pages: list[int] = []

    for page_number in range(1, converted.page_count + 1):
        page = pages_by_number.get(
            page_number,
            ConvertedPage(
                page_number=page_number,
                markdown=f"<!-- Page {page_number}: no extractable text -->",
                needs_ocr=True,
            ),
        )
        if page_number not in converted.pages_needing_ocr:
            updated_pages.append(page)
            continue

        ocr_markdown = ocr_pdf_pages(
            pdf_path,
            pages=[page_number],
            language=config.ocr_language,
            dpi=config.ocr_dpi,
        )
        if ocr_markdown:
            base_markdown = "" if "no extractable text" in page.markdown else page.markdown
            markdown = f"{base_markdown}\n\n## OCR fallback\n\n{ocr_markdown}".strip()
            updated_pages.append(
                replace(
                    page,
                    markdown=markdown,
                    needs_ocr=False,
                    metadata={
                        **page.metadata,
                        "ocrApplied": True,
                        "ocrLanguage": config.ocr_language,
                        "ocrDpi": config.ocr_dpi,
                    },
                )
            )
        else:
            unresolved_pages.append(page_number)
            updated_pages.append(page)

    raw = {
        **converted.raw,
        "ocrApplied": True,
        "ocrLanguage": config.ocr_language,
        "ocrDpi": config.ocr_dpi,
        "originalPdfType": converted.detected_type,
    }
    return replace(
        converted,
        detected_type=f"{converted.detected_type}+OCR",
        pages=updated_pages,
        pages_needing_ocr=unresolved_pages,
        raw=raw,
    )


def convert_pdf(config: LabConfig, file_path: Path) -> ConvertedDocument:
    inspection = inspect_pdf(config.pdf_inspector_command, file_path)
    reader = PdfReader(str(file_path))
    pages: list[ConvertedPage] = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_needs_ocr = page_number in inspection.pages_needing_ocr or not text.strip()
        pages.append(
            ConvertedPage(
                page_number=page_number,
                markdown=_page_to_markdown(page_number, text),
                needs_ocr=page_needs_ocr,
                has_encoding_issues=inspection.has_encoding_issues,
                is_complex_layout=inspection.is_complex_layout,
                metadata={"source": "pypdf", "pdf_inspector": inspection.raw},
            )
        )

    converted = ConvertedDocument(
        source_format="pdf",
        detected_type=inspection.pdf_type,
        confidence=inspection.confidence,
        page_count=inspection.page_count or len(pages),
        pages=pages,
        pages_needing_ocr=inspection.pages_needing_ocr,
        has_encoding_issues=inspection.has_encoding_issues,
        is_complex_layout=inspection.is_complex_layout,
        raw=inspection.raw,
    )

    if should_require_ocr(inspection):
        return _apply_pdf_ocr_fallback(config, file_path, converted)
    return converted


def _run_has_page_break(run) -> bool:
    for br in run._element.xpath(".//w:br"):
        if br.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
            return True
    return False


def _paragraph_to_markdown(paragraph) -> tuple[str, bool]:
    text = _clean_text(paragraph.text)
    has_page_break = any(_run_has_page_break(run) for run in paragraph.runs)
    if not text:
        return "", has_page_break

    style = (paragraph.style.name if paragraph.style is not None else "").lower()
    if style.startswith("heading"):
        level_text = "".join(ch for ch in style if ch.isdigit())
        level = min(max(int(level_text or "1"), 1), 6)
        return f"{'#' * level} {text}", has_page_break
    if "list" in style:
        return f"- {text}", has_page_break
    return text, has_page_break


def _table_to_markdown(table) -> str:
    rows = [[_clean_text(cell.text).replace("\n", " ") for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _markdown_pages_from_blocks(blocks: list[str], *, source_format: str) -> list[ConvertedPage]:
    pages: list[ConvertedPage] = []
    current: list[str] = []

    def flush() -> None:
        if not current:
            return
        page_number = len(pages) + 1
        pages.append(
            ConvertedPage(
                page_number=page_number,
                markdown=f"<!-- Logical Page {page_number} ({source_format.upper()}) -->\n\n"
                + "\n\n".join(current).strip(),
                metadata={"source": source_format, "page_kind": "logical"},
            )
        )
        current.clear()

    for block in blocks:
        if block == "\f":
            flush()
            continue
        if block.strip():
            current.append(block.strip())
    flush()

    if not pages:
        pages.append(
            ConvertedPage(
                page_number=1,
                markdown=f"<!-- Logical Page 1 ({source_format.upper()}): no extractable text -->",
                metadata={"source": source_format, "page_kind": "logical"},
            )
        )
    return pages


def convert_docx(file_path: Path) -> ConvertedDocument:
    document = Document(str(file_path))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        markdown, has_page_break = _paragraph_to_markdown(paragraph)
        if markdown:
            blocks.append(markdown)
        if has_page_break:
            blocks.append("\f")

    for table in document.tables:
        markdown = _table_to_markdown(table)
        if markdown:
            blocks.append(markdown)

    pages = _markdown_pages_from_blocks(blocks, source_format="docx")
    return ConvertedDocument(
        source_format="docx",
        detected_type="DocxText",
        confidence=0.95 if any(page.markdown.strip() for page in pages) else 0.0,
        page_count=len(pages),
        pages=pages,
        raw={"engine": "python-docx"},
    )


def _textutil_binary() -> str:
    binary = shutil.which("textutil") or "/usr/bin/textutil"
    if not Path(binary).exists():
        raise RuntimeError("textutil is required to convert DOC/RTF files in the local lab.")
    return binary


def convert_textutil_document(file_path: Path, source_format: str) -> ConvertedDocument:
    with tempfile.TemporaryDirectory() as tmp:
        output_path = Path(tmp) / "document.txt"
        subprocess.run(
            [_textutil_binary(), "-convert", "txt", "-output", str(output_path), str(file_path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=120,
        )
        text = output_path.read_text(encoding="utf-8", errors="replace")

    blocks = [_clean_text(block) for block in text.split("\f")]
    page_blocks: list[str] = []
    for index, block in enumerate(blocks):
        if index > 0:
            page_blocks.append("\f")
        page_blocks.append(block)

    pages = _markdown_pages_from_blocks(page_blocks, source_format=source_format)
    return ConvertedDocument(
        source_format=source_format,
        detected_type=f"{source_format.upper()}Text",
        confidence=0.9 if text.strip() else 0.0,
        page_count=len(pages),
        pages=pages,
        raw={"engine": "textutil"},
    )


def convert_document(config: LabConfig, file_path: Path, source_format: str) -> ConvertedDocument:
    if source_format == "pdf":
        return convert_pdf(config, file_path)
    if source_format == "docx":
        return convert_docx(file_path)
    if source_format in {"doc", "rtf"}:
        return convert_textutil_document(file_path, source_format)
    raise ValueError(f"Unsupported lab source format: {source_format}")
