"""SDD 008 T024 tests: matriz DOCX renderer."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.matriz_docx_renderer import (
    UnknownDocxNodeError,
    UnresolvedTokenError,
    render_minuta_docx,
)


def _docx(bytes_: bytes) -> Document:
    return Document(BytesIO(bytes_))


def test_render_minuta_docx_preserves_order_and_title_block_verbatim():
    primero = "PRIMERO: El inmueble se encuentra inscrito a fojas ciento veinte."
    docx_bytes = render_minuta_docx(
        metadata={"title_lines": ["Minuta de compraventa"]},
        clauses=[
            {
                "clause_key": "comparecencia",
                "title": "COMPARECENCIA",
                "resolved_content": {
                    "schema_version": 1,
                    "type": "doc",
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [{"type": "text", "text": "Comparecen las partes."}],
                        }
                    ],
                },
            },
            {
                "clause_key": "antecedentes_dominio",
                "title": "ANTECEDENTES DE DOMINIO DEL PREDIO MATRIZ",
                "resolved_content": {
                    "schema_version": 1,
                    "type": "doc",
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [{"type": "text", "text": primero}],
                        }
                    ],
                },
            },
            {
                "clause_key": "precio_liquidacion",
                "title": "PRECIO Y LIQUIDACIÓN",
                "resolved_content": {
                    "schema_version": 1,
                    "type": "doc",
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [{"type": "text", "text": "El precio se paga al contado."}],
                        }
                    ],
                },
            },
            {
                "clause_key": "entrega_material",
                "title": "ENTREGA MATERIAL",
                "resolved_content": {
                    "schema_version": 1,
                    "type": "doc",
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [{"type": "text", "text": "La entrega se realiza en este acto."}],
                        }
                    ],
                },
            },
        ],
    )
    document = _docx(docx_bytes)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    assert paragraphs[0] == "MINUTA DE COMPRAVENTA"
    assert paragraphs[1].startswith("COMPARECENCIA. Comparecen")
    assert paragraphs[2] == primero
    assert paragraphs[3].startswith("SEGUNDO: PRECIO Y LIQUIDACIÓN. El precio")
    assert paragraphs[4].startswith("TERCERO: ENTREGA MATERIAL. La entrega")
    assert document.paragraphs[1].runs[0].bold is True
    assert document.paragraphs[2].runs[0].bold is True
    assert document.paragraphs[3].runs[0].bold is True
    assert document.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_render_minuta_docx_rejects_unresolved_tokens():
    with pytest.raises(UnresolvedTokenError) as exc:
        render_minuta_docx(
            clauses=[
                {
                    "clause_key": "precio",
                    "title": "PRECIO",
                    "resolved_content": {
                        "schema_version": 1,
                        "type": "doc",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "variable_token",
                                        "attrs": {"variableKey": "transaccion.precio"},
                                    }
                                ],
                            }
                        ],
                    },
                }
            ]
        )

    assert exc.value.variable_keys == ["transaccion.precio"]


def test_render_minuta_docx_rejects_unknown_nodes():
    with pytest.raises(UnknownDocxNodeError) as exc:
        render_minuta_docx(
            clauses=[
                {
                    "clause_key": "x",
                    "title": "X",
                    "resolved_content": {
                        "schema_version": 1,
                        "type": "doc",
                        "content": [{"type": "table", "content": []}],
                    },
                }
            ]
        )

    assert exc.value.node_types == ["table"]


# ─── SDD 011 T015: marca de borrador (FR-008 / ADR-009) ──────────────────────

_NOTICE_CLAUSE = {
    "clause_key": "comparecencia",
    "title": "COMPARECENCIA",
    "resolved_content": {
        "schema_version": 1,
        "type": "doc",
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "Comparecen."}]}
        ],
    },
}


def test_render_minuta_docx_includes_draft_notice_when_requested():
    """La marca de borrador aparece visible en el DOCX entregable (FR-008)."""
    notice = "Borrador sujeto a revisión legal"
    docx_bytes = render_minuta_docx(
        metadata={"title_lines": ["Minuta de compraventa"], "draft_notice": notice},
        clauses=[_NOTICE_CLAUSE],
    )
    texts = [paragraph.text for paragraph in _docx(docx_bytes).paragraphs]
    assert notice.upper() in texts


def test_render_minuta_docx_omits_draft_notice_by_default():
    """Sin draft_notice el DOCX no cambia: no-regresion del renderer."""
    docx_bytes = render_minuta_docx(
        metadata={"title_lines": ["Minuta de compraventa"]},
        clauses=[_NOTICE_CLAUSE],
    )
    texts = [paragraph.text for paragraph in _docx(docx_bytes).paragraphs]
    assert not any("BORRADOR SUJETO" in text.upper() for text in texts)
