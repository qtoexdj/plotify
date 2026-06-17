"""
DEPRECATED: generador MVP de documentos legales en PDF/DOCX.

SDD 008 usa `matriz_docx_renderer.py` + `escritura_minuta_generations` para
minutas de escritura. Este modulo queda solo para compatibilidad del router
legacy `/documents/*`; no debe conectarse a rutas nuevas.

Responsabilidades:
    - generate_pdf: renderiza template → HTML → PDF usando WeasyPrint.
    - generate_docx: renderiza template → texto plano → DOCX usando python-docx.
    - persist_document: sube archivo a Supabase Storage y registra en generated_documents.

Ref: plotify_memori/Generacion de Documentos.md
"""

import asyncio
import time
from io import BytesIO
from typing import Any, TypedDict

from weasyprint import HTML  # type: ignore
from docx import Document as DocxDocument  # type: ignore
from docx.shared import Pt  # type: ignore

from services.document_engine import render_template, resolve_variables
from core.database import get_supabase_client
from core.logger import get_logger

logger = get_logger(__name__)

# CSS base para documentos legales chilenos (tamaño carta, márgenes notariales)
CSS_TEMPLATE = """
@page {
    size: letter;
    margin: 25mm;
}
body {
    font-family: 'Times New Roman', serif;
    font-size: 12pt;
    line-height: 1.6;
    text-align: justify;
}
h1 { font-size: 16pt; text-align: center; margin-bottom: 12pt; }
h2 { font-size: 14pt; margin-top: 12pt; }
p  { margin-bottom: 8pt; }
.firma {
    margin-top: 60px;
    display: flex;
    justify-content: space-around;
}
.firma-line {
    border-top: 1px solid black;
    width: 200px;
    text-align: center;
    padding-top: 5px;
    font-size: 10pt;
}
"""


class PersistedDocument(TypedDict):
    id: str
    file_url: str
    file_format: str
    document_type: str
    version_number: int
    lot_id: str
    template_id: str
    missing_variables_accepted: bool
    missing_variables: list[str]
    selected_recipients: list[str]
    delivery_status: str
    delivery_failed_attempts: int
    delivery_error_message: str | None


async def generate_pdf(template_id: str, lot_id: str, organization_id: str) -> bytes:
    """
    Genera un PDF a partir de un template y los datos del lote.

    Returns:
        Bytes del PDF generado.
    """
    html_content = await render_template(template_id, lot_id, organization_id)
    full_html = (
        f"<html><head>"
        f"<meta charset='utf-8'>"
        f"<style>{CSS_TEMPLATE}</style>"
        f"</head><body>{html_content}</body></html>"
    )
    # WeasyPrint es síncrono — ejecutar en thread para no bloquear el event loop
    pdf_bytes: bytes = await asyncio.to_thread(
        lambda: HTML(string=full_html).write_pdf()
    )
    return pdf_bytes


async def generate_docx(template_id: str, lot_id: str, organization_id: str) -> bytes:
    """
    Genera un DOCX a partir de un template y los datos del lote.

    Convierte el HTML renderizado a párrafos de texto limpio en un documento Word.

    Returns:
        Bytes del DOCX generado.
    """
    html_content = await render_template(template_id, lot_id, organization_id)

    def _build_docx() -> bytes:
        doc = DocxDocument()
        # Configurar fuente base
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        # Parseo simple: separar por líneas; líneas con <h1>/<h2> → títulos
        import re

        lines = html_content.split("\n")
        for line in lines:
            clean = re.sub(r"<[^>]+>", "", line).strip()
            if not clean:
                continue
            if re.search(r"<h1", line, re.IGNORECASE):
                doc.add_heading(clean, level=1)
            elif re.search(r"<h2", line, re.IGNORECASE):
                doc.add_heading(clean, level=2)
            else:
                doc.add_paragraph(clean)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    docx_bytes: bytes = await asyncio.to_thread(_build_docx)
    return docx_bytes


async def persist_document(
    *,
    file_bytes: bytes,
    file_format: str,
    template_id: str,
    lot_id: str,
    organization_id: str,
    generated_by: str | None = None,
    document_type: str = "generated",
    missing_variables_accepted: bool = False,
    missing_variables: list[str] | None = None,
    selected_recipients: list[str] | None = None,
) -> PersistedDocument:
    """
    Sube el documento a Supabase Storage y registra en generated_documents.

    Returns:
        Metadata persistida del documento generado.
    """
    supabase = get_supabase_client()
    content_type_map = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    content_type = content_type_map.get(file_format, "application/octet-stream")
    file_path = f"{organization_id}/docs/{lot_id}_{int(time.time())}.{file_format}"

    await asyncio.to_thread(
        lambda: supabase.storage.from_("documents").upload(
            file_path,
            file_bytes,
            {"content-type": content_type},
        )
    )

    # Obtener URL (bucket privado → signed URL con 7 días de expiración)
    signed = await asyncio.to_thread(
        lambda: supabase.storage.from_("documents").create_signed_url(
            file_path, expires_in=604800
        )
    )
    file_url: str = signed.get("signedURL") or signed.get("signedUrl", file_path)

    # Snapshot de variables para trazabilidad legal (inmutable)
    variables = await resolve_variables(lot_id, organization_id)
    missing_variables = missing_variables or []
    selected_recipients = selected_recipients or []

    latest_version_result = await asyncio.to_thread(
        lambda: (
            supabase.table("generated_documents")
            .select("version_number")
            .eq("lot_id", lot_id)
            .eq("template_id", template_id)
            .eq("document_type", document_type)
            .order("version_number", desc=True)
            .limit(1)
            .execute()
        )
    )
    latest_rows: list[dict[str, Any]] = latest_version_result.data or []
    latest_version = (
        int(latest_rows[0].get("version_number") or 0) if latest_rows else 0
    )
    version_number = latest_version + 1

    record: dict = {
        "organization_id": organization_id,
        "template_id": template_id,
        "lot_id": lot_id,
        "document_type": document_type,
        "file_url": file_url,
        "file_format": file_format,
        "variables_snapshot": variables,
        "version_number": version_number,
        "missing_variables_accepted": missing_variables_accepted,
        "missing_variables": missing_variables,
        "selected_recipients": selected_recipients,
        "delivery_status": "pending",
        "delivery_failed_attempts": 0,
    }
    if generated_by:
        record["generated_by"] = generated_by

    insert_result = await asyncio.to_thread(
        lambda: supabase.table("generated_documents").insert(record).execute()
    )
    inserted = (insert_result.data or [{}])[0]

    logger.info(
        "Documento generado y persistido",
        file_path=file_path,
        format=file_format,
        lot_id=lot_id,
        organization_id=organization_id,
    )
    return {
        "id": inserted.get("id", ""),
        "file_url": inserted.get("file_url", file_url),
        "file_format": inserted.get("file_format", file_format),
        "document_type": inserted.get("document_type", document_type),
        "version_number": int(inserted.get("version_number") or version_number),
        "lot_id": inserted.get("lot_id", lot_id),
        "template_id": inserted.get("template_id", template_id),
        "missing_variables_accepted": bool(
            inserted.get("missing_variables_accepted", missing_variables_accepted)
        ),
        "missing_variables": inserted.get("missing_variables", missing_variables),
        "selected_recipients": inserted.get(
            "selected_recipients", selected_recipients
        ),
        "delivery_status": inserted.get("delivery_status", "pending"),
        "delivery_failed_attempts": int(
            inserted.get("delivery_failed_attempts") or 0
        ),
        "delivery_error_message": inserted.get("delivery_error_message"),
    }
