"""SDD 008: ProseMirror JSON -> DOCX renderer (research D5).

Receives clause content ALREADY resolved by ``matriz_token_resolution``
(no token lookups here) and renders the minuta with python-docx: clause
titles in bold running text, legal ordinal numbering, justified paragraphs,
no tables. Unknown node types raise :class:`UnknownDocxNodeError` listing
the offending nodes — silent omission is forbidden (spec edge case).

Skeleton (T011): node walker + explicit unknown-node error; full minuta
styling lands in T024.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from core.logger import get_logger

logger = get_logger(__name__)

KNOWN_BLOCK_NODE_TYPES = frozenset(
    ("paragraph", "block_token", "repeat_section", "conditional_section")
)
KNOWN_INLINE_NODE_TYPES = frozenset(("text", "variable_token"))
KNOWN_NODE_TYPES = KNOWN_BLOCK_NODE_TYPES | KNOWN_INLINE_NODE_TYPES | {"doc"}
LEGAL_ORDINALS = (
    "PRIMERO",
    "SEGUNDO",
    "TERCERO",
    "CUARTO",
    "QUINTO",
    "SEXTO",
    "SÉPTIMO",
    "OCTAVO",
    "NOVENO",
    "DÉCIMO",
    "UNDÉCIMO",
    "DUODÉCIMO",
    "DÉCIMO TERCERO",
    "DÉCIMO CUARTO",
    "DÉCIMO QUINTO",
    "DÉCIMO SEXTO",
    "DÉCIMO SÉPTIMO",
    "DÉCIMO OCTAVO",
    "DÉCIMO NOVENO",
    "VIGÉSIMO",
)
_LEGAL_ORDINAL_RE = re.compile(
    r"^\s*("
    + "|".join(re.escape(ordinal) for ordinal in LEGAL_ORDINALS)
    + r")\s*:",
    re.IGNORECASE,
)


class MatrizDocxError(Exception):
    """Base error for DOCX rendering failures."""


class UnknownDocxNodeError(MatrizDocxError):
    """Unknown node types in resolved content: export must fail explicitly."""

    def __init__(self, node_types: list[str]):
        self.node_types = sorted(set(node_types))
        super().__init__(
            "Cannot render unknown ProseMirror node types to DOCX: "
            + ", ".join(self.node_types)
        )


class UnresolvedTokenError(MatrizDocxError):
    """A variable_token survived resolution: generation must abort (SC-002)."""

    def __init__(self, variable_keys: list[str]):
        self.variable_keys = sorted(set(variable_keys))
        super().__init__(
            "Resolved matriz content still contains unresolved tokens: "
            + ", ".join(self.variable_keys)
        )


def collect_unknown_node_types(node: dict[str, Any]) -> list[str]:
    """Walk a (resolved) ProseMirror document collecting unknown node types."""
    unknown: list[str] = []
    node_type = node.get("type")
    if node_type not in KNOWN_NODE_TYPES:
        unknown.append(str(node_type))
    for child in node.get("content") or []:
        if isinstance(child, dict):
            unknown.extend(collect_unknown_node_types(child))
    return unknown


def collect_unresolved_token_keys(node: dict[str, Any]) -> list[str]:
    """Resolved content must carry zero variable_token nodes (SC-002)."""
    keys: list[str] = []
    if node.get("type") == "variable_token":
        keys.append(str((node.get("attrs") or {}).get("variableKey")))
    for child in node.get("content") or []:
        if isinstance(child, dict):
            keys.extend(collect_unresolved_token_keys(child))
    return keys


def render_minuta_docx(
    *,
    clauses: list[dict[str, Any]],
    metadata: dict[str, Any] | None = None,
) -> bytes:
    """Render resolved clauses to a DOCX binary.

    Input is the resolved clause list from ``matriz_token_resolution``. The
    renderer performs no lookups and refuses any unresolved token remnants.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:  # pragma: no cover - dependency is pinned.
        raise MatrizDocxError("python-docx is required to render matriz DOCX") from exc

    unknown_nodes: list[str] = []
    unresolved_tokens: list[str] = []
    for clause in clauses:
        content = clause.get("resolved_content") or clause.get("content_json") or {}
        unknown_nodes.extend(collect_unknown_node_types(content))
        unresolved_tokens.extend(collect_unresolved_token_keys(content))
    if unknown_nodes:
        raise UnknownDocxNodeError(unknown_nodes)
    if unresolved_tokens:
        raise UnresolvedTokenError(unresolved_tokens)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    if metadata:
        title = str(metadata.get("title") or "Minuta de compraventa")
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(title.upper())
        run.bold = True

        # SDD 011 (FR-008 / ADR-009): marca visible de borrador. Opt-in via
        # metadata (el endpoint la pasa siempre); ausente => DOCX sin cambios
        # (no-regresion del renderer).
        draft_notice = metadata.get("draft_notice")
        if draft_notice:
            notice = document.add_paragraph()
            notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
            notice_run = notice.add_run(str(draft_notice).upper())
            notice_run.bold = True
            notice_run.italic = True

    next_ordinal_index = 0
    for clause in clauses:
        content = clause.get("resolved_content") or clause.get("content_json") or {}
        title = str(clause.get("title") or clause.get("clause_key") or "").strip()
        clause_key = str(clause.get("clause_key") or "")
        paragraphs = _content_paragraph_texts(content)
        if not paragraphs:
            continue

        title_consumed = False
        first_ordinal = _legal_ordinal_prefix(paragraphs[0])
        if first_ordinal is not None:
            next_ordinal_index = max(next_ordinal_index, first_ordinal + 1)

        for index, text in enumerate(paragraphs):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if index == 0 and first_ordinal is not None:
                label_end = text.index(":") + 1
                paragraph.add_run(text[:label_end]).bold = True
                paragraph.add_run(text[label_end:])
                title_consumed = True
                continue
            if (
                index == 0
                and title
                and not _starts_with_title(text, title)
                and _uses_legal_ordinal(clause_key=clause_key)
            ):
                ordinal = _ordinal_at(next_ordinal_index)
                next_ordinal_index += 1
                title_run = paragraph.add_run(f"{ordinal}: {title}. ")
                title_run.bold = True
                title_consumed = True
            elif index == 0 and title and not _starts_with_title(text, title):
                title_run = paragraph.add_run(f"{title}. ")
                title_run.bold = True
                title_consumed = True
            paragraph.add_run(text)
        if title and not title_consumed and paragraphs:
            # Existing legal text already carries the ordinal/title; make the
            # leading label bold without changing the underlying words.
            first = document.paragraphs[-len(paragraphs)]
            text = first.text
            if text.upper().startswith(title.upper()):
                first.clear()
                first.add_run(text[: len(title)]).bold = True
                first.add_run(text[len(title) :])

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _content_paragraph_texts(content: dict[str, Any]) -> list[str]:
    texts: list[str] = []
    for node in content.get("content") or []:
        if not isinstance(node, dict):
            continue
        if node.get("type") == "paragraph":
            text = _inline_text(node.get("content") or [])
            if text.strip():
                texts.append(text)
        elif node.get("type") in {"repeat_section", "conditional_section"}:
            texts.extend(_content_paragraph_texts({"content": node.get("content") or []}))
    return texts


def _inline_text(nodes: list[Any]) -> str:
    parts: list[str] = []
    for node in nodes:
        if isinstance(node, dict) and node.get("type") == "text":
            parts.append(str(node.get("text") or ""))
    return "".join(parts)


def _starts_with_title(text: str, title: str) -> bool:
    normalized_text = text.strip().upper()
    normalized_title = title.strip().upper()
    return normalized_text.startswith(normalized_title)


def _legal_ordinal_prefix(text: str) -> int | None:
    match = _LEGAL_ORDINAL_RE.match(text)
    if not match:
        return None
    normalized = match.group(1).upper()
    return LEGAL_ORDINALS.index(normalized) if normalized in LEGAL_ORDINALS else None


def _ordinal_at(index: int) -> str:
    if index >= len(LEGAL_ORDINALS):
        raise MatrizDocxError("Too many clauses for built-in legal ordinal numbering")
    return LEGAL_ORDINALS[index]


def _uses_legal_ordinal(*, clause_key: str) -> bool:
    return clause_key != "comparecencia"
